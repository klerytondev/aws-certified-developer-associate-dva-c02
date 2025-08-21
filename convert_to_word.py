#!/usr/bin/env python3
"""
Script para converter arquivo Markdown para Word usando python-docx e markdown
"""

import markdown
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

def markdown_to_word(markdown_file, output_file):
    """Converte arquivo Markdown para Word"""
    
    # Ler o arquivo markdown
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Criar documento Word
    doc = Document()
    
    # Processar linha por linha
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Linha vazia - adicionar parágrafo vazio
            doc.add_paragraph()
            continue
            
        # Títulos
        if line.startswith('# '):
            title = line[2:].strip()
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        elif line.startswith('## '):
            subtitle = line[3:].strip()
            doc.add_heading(subtitle, level=2)
            
        elif line.startswith('### '):
            subtitle = line[4:].strip()
            doc.add_heading(subtitle, level=3)
            
        # Divisão com imagem centralizada
        elif '<div align="center">' in line:
            continue  # Pular a linha da div
            
        elif '<img src="image.png"' in line:
            # Adicionar texto indicando onde está a imagem
            p = doc.add_paragraph('[IMAGEM: AWS Certified Developer Associate DVA-C02]')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        elif '</div>' in line:
            continue  # Pular fechamento da div
            
        # Texto em negrito com **
        elif line.startswith('**') and line.endswith('**'):
            text = line[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            
        # Listas com marcadores
        elif line.startswith('- '):
            item_text = line[2:].strip()
            # Remover formatação markdown básica
            item_text = re.sub(r'\*\*(.*?)\*\*', r'\1', item_text)  # Negrito
            doc.add_paragraph(item_text, style='List Bullet')
            
        # Listas numeradas  
        elif re.match(r'^\d+\.', line):
            item_text = re.sub(r'^\d+\.\s*', '', line)
            item_text = re.sub(r'\*\*(.*?)\*\*', r'\1', item_text)
            doc.add_paragraph(item_text, style='List Number')
            
        # Links
        elif line.startswith('**[') and '](' in line:
            # Extrair texto e URL do link
            match = re.match(r'\*\*\[(.*?)\]\((.*?)\)\*\*', line)
            if match:
                text, url = match.groups()
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                p.add_run(f' ({url})')
                
        # Tabelas (formato markdown simples)
        elif '|' in line and not line.startswith('|---'):
            # Isso é parte de uma tabela - para simplicidade, vamos adicionar como texto
            table_text = line.replace('|', ' | ')
            doc.add_paragraph(table_text)
            
        # Texto com emojis/checkmarks
        elif line.startswith('✅'):
            doc.add_paragraph(line, style='List Bullet')
            
        # Linha de separação
        elif line.startswith('---'):
            doc.add_paragraph('_' * 50)
            
        # Texto normal
        elif line:
            # Remover formatação markdown básica
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)  # Negrito
            clean_text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', clean_text)  # Links
            
            p = doc.add_paragraph(clean_text)
    
    # Salvar documento
    doc.save(output_file)
    print(f"Arquivo Word criado com sucesso: {output_file}")

if __name__ == "__main__":
    markdown_file = "resumo-executivo-dva-c02.md"
    output_file = "resumo-executivo-dva-c02.docx"
    
    if os.path.exists(markdown_file):
        markdown_to_word(markdown_file, output_file)
    else:
        print(f"Arquivo {markdown_file} não encontrado!")
